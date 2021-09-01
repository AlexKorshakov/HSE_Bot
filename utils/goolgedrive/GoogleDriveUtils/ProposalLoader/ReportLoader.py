import os
from pprint import pprint
from time import sleep

from docx import Document
from tqdm import tqdm

from Config.Config import PATH_TO_JSON, SERVICE_ACCOUNT_FILE, SPREAD_SHEET_ID
from GoogleSheetsWorker.GoogleSheetsWorker import account_credentials, get_data_from_google
from Json import JsonWork as js

__date__ = '05.05.2021'
print(__name__ + " *** " + __date__)
print(f'Invoking __init__.py for {__name__}')

MAIN_FOLDER: str = r"C:\Users\DeusEx\Desktop\!Ликвидатор\ЦОТЭБ\Ликвидатор"
TRIGGER: str = r"Выписки"
EXCLUSIVE_TRIGGER: str = r"закрыта"

EXTENSION_DOCX = ".docx"


def load_table_from_docx(word_doc) -> list:
    """ Чтение и получение данных из файлов .docx (files_with_data)
    :return:
    """
    table_text = []

    for row_number, row in enumerate(word_doc.tables[0].rows):
        if row_number < 2:
            continue
        table_text.append([cell.text for cell in row.cells])
    return table_text


def load_paragraphs_from_docx(word_doc) -> list:
    """ Чтение и получение данных из файлов .docx (files_with_data)
    :return:
    """
    paragraphs_text = []

    with open('../ExeptionList.txt', 'r', encoding="UTF-8") as file:
        exception_list = [line.rstrip() for line in set(file) if line != '' or line != ' ']

    for row_number, row in enumerate(word_doc.paragraphs):
        if row_number < 2:
            continue
        if row.text not in exception_list and row.text != "":
            paragraphs_text.append(row.text)

    js.write_json_file(data=paragraphs_text, name=PATH_TO_JSON + "paragraphs_text.json")

    return paragraphs_text


def _read_files(file_with_data: str, file_extension: str):
    """Читаем данные из файла с учетом расширения
    :type file_extension: str : расширение файлам
    :return:
    """
    if not _check_is_file(file=file_with_data):
        return

    if file_extension != EXTENSION_DOCX:
        return
    word_doc = Document(file_with_data)

    table_text = load_table_from_docx(word_doc)
    paragraphs_text = load_paragraphs_from_docx(word_doc)
    return table_text, paragraphs_text


def _create_list_for_load(data_for_load) -> list:
    """Формирование / разворачивание списка list из List[list, list, ...] в List[..., ...]
    :param data_for_load: список
    :return: list
    """
    return [i for item in data_for_load for i in item]


def _find_training_program(program):
    """ Поиск и сопоставление в заявке программам

    :param program:
    :return:
    """
    if program is None:
        return ""

    search_phrase = ['Ответственный исполнитель работ на высоте 2 группа по безопасности работ на высоте']
    if program in search_phrase:
        training_program = "Обучение безопасным методам и приемам выполнения работ на высоте"
    else:
        training_program = program

    return training_program.rstrip(" ")


def _find_user(item, range_workers, report) -> object:
    """Поиск работника в безе
    :rtype: object
    """

    f_user = item[1].rstrip(" ") if item[0].isnumeric() else item[0].rstrip(" ")
    training_program = _find_training_program(program=item[6])

    for user in range_workers:
        try:
            user_from_tab = user[4]
        except IndexError:
            continue

        if f_user != user_from_tab or item[1] == '':
            continue
        # print(f'{f_user} find in {user[4]} number: {user[1]} training_program: {training_program} !!!')

        return {'training_program': training_program,
                'user': user[4],
                'number': user[1],
                'reports': report
                }
    pprint(f"User not found {f_user} ")
    return None


def _collect_data(*, reports:dict, range_workers) -> list:
    """Сбоор данных из ранее прочитанных файлов
    :type reports: dict
    """
    global user_data
    assert reports is not None, "data not found!!"

    collect_data = []

    print(f'\n{"*" * 15} reports: {reports["file"]} {"*" * 15}')

    for item in reports["data"]:
        if item[0] != '':
            user_data = _find_user(item=item, range_workers=range_workers, report=reports["file"])

        if user_data is None:
            continue
        collect_data.append(user_data)

    return collect_data


def _check_is_file(*, file) -> bool:
    """Проверка наличия файла
    """
    return os.path.isfile(file)


def _check_is_folder(path_folder) -> bool:
    """Проверяем существует ли папка. Если не существует - создаём.
    """
    return True if os.path.exists(path_folder) else False


def _collect_folder() -> list:
    """Формирование списка list файлов .pdf  в целевой папке MAIN_FOLDER
    :return: list files_with_data
    """
    folders = os.listdir(MAIN_FOLDER)  # просматриваем папку

    return [x for x in folders
            if _check_is_folder(MAIN_FOLDER + "\\" + x) and
            TRIGGER in x and
            EXCLUSIVE_TRIGGER not in x]


def _collect_file(folder, report_extension) -> list:
    """Формирование списка list файлов .pdf  в целевой папке MAIN_FOLDER
    :return: list files_with_data
    """
    try:
        if _check_is_file(file=MAIN_FOLDER + "\\" + folder):
            return [folder]

        report_files = os.listdir(MAIN_FOLDER + "\\" + folder)

        # просматриваем папку
    except (FileNotFoundError, NotADirectoryError):
        return []

    return [x for x in report_files if x.endswith(report_extension) and x[0] != "~"]


def load_reports(service, reports_files: object, reports_extension: str, range_workers: list):
    """Загрузка файлов заявок из папки
    :param reports_extension: расширение файлов
    :param reports_files:
    :return:
    """
    if not isinstance(reports_files, list):
        assert reports_files is not None, "reports_files not found!!"

    data_for_load = []

    p_bar = tqdm(reports_files, total=len(reports_files), leave=True, colour='green', dynamic_ncols=True, disable=False)

    for item_index, item in enumerate(p_bar, start=1):
        p_bar.set_description(f"Progress (file {item_index})")
        p_bar.refresh()  # to show immediately the update
        p_bar.set_postfix({'file': item["file"]})
        sleep(0.1)

        data, paragraphs_text = _read_files(item["path"], reports_extension)

        program = _define_training_program(paragraphs_text=paragraphs_text)

        datas = []
        for worker_data in data:
            worker_data.append(program)
            datas.append(worker_data)

        data_from_files: dict = {"file": item['file'].replace(reports_extension, ''),
                                 "data": datas,
                                 "program": program}
        data = _collect_data(reports=data_from_files, range_workers=range_workers)
        data_for_load.append(data)

    data_for_load = _create_list_for_load(data_for_load)
    pprint(data_for_load)

    options = {"valueRenderOption": 'UNFORMATTED_VALUE',
               "dateTimeRenderOption": 'FORMATTED_STRING'}

    range_data_id = get_data_from_google(service=service,
                                         spread_sheet_id=SPREAD_SHEET_ID,
                                         sheet_range=range_workers,
                                         name="sheet_values_workers",
                                         options=options)

    range_folders = [number_item[1] for number_item in range_data_id]

    _load_report(service=service, data_for_load=_create_list_for_load(data_for_load),
                 range_folders=range_folders)


def _define_date(paragraphs_text: list) -> str:
    pass


def _define_training_program(paragraphs_text: list) -> str:
    """
    :param paragraphs_text: list
    :return:
    """

    training_list: list = js.read_json_file(file=PATH_TO_JSON + "trainings_programs.json")
    if not isinstance(training_list, list):
        assert training_list is not None, "trainings_programs.json not found!!"

    for item in paragraphs_text:
        for program in training_list:

            if item == program.get('juxtaposition'):
                return program.get('program')
    return


def _post_value(service, spread_sheet_id, sheet_id, row_index, column_index, text):
    """
    :param service:
    :param spread_sheet_id:
    :param sheet_id:
    :param row_index:
    :param column_index:
    :param text:
    :return:
    """
    requests = [{"updateCells": {
        "rows": [
            {
                "values": [{
                    "userEnteredValue": {
                        "stringValue": text
                    }
                }]
            }
        ],
        "fields": "userEnteredValue",
        "start": {
            "sheetId": sheet_id,
            "rowIndex": row_index,
            "columnIndex": column_index
        }
    }}]
    body = {"requests": requests}
    request = service.spreadsheets().batchUpdate(spreadsheetId=spread_sheet_id, body=body).execute()
    return request


def _get_column_index(training_program) -> object:
    programs = js.read_json_file(file=PATH_TO_JSON + "trainings_programs.json")

    range_programs = [program['program'] for program in programs]

    for _ in programs:

        if training_program in range_programs:
            ind = range_programs.index(training_program)
            var = programs[ind]['index']
            return int(var)
        else:
            print("program not found!!!")
            return None


def _read_data_from_files(reports_files: list, reports_extension: str, range_workers) -> list:
    """
    :param reports_files:
    :param reports_extension:
    :return:
    """
    if not isinstance(reports_files, list):
        assert reports_files is not None, "reports_files not found!!"

    data_for_load = []

    p_bar = tqdm(reports_files, total=len(reports_files), leave=True, colour='green', dynamic_ncols=True, disable=False)

    for item_index, item in enumerate(p_bar, start=1):
        p_bar.set_description(f"Progress (file {item_index})")
        p_bar.refresh()  # to show immediately the update
        p_bar.set_postfix({'file': item["file"]})
        sleep(0.1)

        data, paragraphs_text = _read_files(item["path"], reports_extension)

        report_program = _define_training_program(paragraphs_text=paragraphs_text)
        report_start_date = _define_date(paragraphs_text=paragraphs_text)
        report_end_date = _define_date(paragraphs_text=paragraphs_text)

        datas = []
        for worker_data in data:
            worker_data.append(report_program)
            datas.append(worker_data)

        data_from_files: dict = {"file": item['file'].replace(reports_extension, ''),
                                 "data": datas,
                                 "program": report_program}
        data = _collect_data(reports=data_from_files, range_workers=range_workers)
        data_for_load.append(data)

    return data_for_load


def _load_reports(*, service, range_workers, data_for_load: list, ):
    """
    :param data_for_load: list
    :return:
    """

    data_for_load = _create_list_for_load(data_for_load)
    js.write_json_file(data=data_for_load, name=PATH_TO_JSON + "report_data_for_load")
    pprint(data_for_load)

    if not service:
        service = account_credentials(service_account_file=SERVICE_ACCOUNT_FILE)

    options = {"valueRenderOption": 'UNFORMATTED_VALUE',
               "dateTimeRenderOption": 'FORMATTED_STRING'}

    range_data_id = get_data_from_google(service=service,
                                         spread_sheet_id=SPREAD_SHEET_ID,
                                         sheet_range=range_workers,
                                         name="sheet_values_workers",
                                         options=options)

    range_folders = [number_item[1] for number_item in range_data_id]

    _load_report(service=service, data_for_load=data_for_load, range_folders=range_folders)


def _load_report(*, service, data_for_load: list, range_folders: list):
    """ Загрузка найденный в папках отчетов
    :param data_for_load:
    :return:
    """
    not_entered_values = []

    p_bar = tqdm(data_for_load, total=len(data_for_load), leave=True, colour='green', dynamic_ncols=True, disable=False)

    for item_index, item in enumerate(p_bar, start=1):
        p_bar.set_description(f"Progress reports (file {item_index})")
        p_bar.refresh()  # to show immediately the update
        p_bar.set_postfix({'reports': item["reports"]})
        sleep(0.1)

        sleep(10) if item_index % 10 == 0 else sleep(0.1)

        if item['number'] in range_folders:

            row_index = range_folders.index(item['number']) + 9 - 1
            text = item['reports']
            column_index = _get_column_index(item['training_program'])
            sheet_id = 0

            if column_index is None:
                print(f'item not found {item["number"]}')
                not_entered_values.append(item)
                continue

            try:
                _post_value(service=service, spread_sheet_id=SPREAD_SHEET_ID, sheet_id=sheet_id,
                            row_index=row_index, column_index=column_index, text=text)
            except Exception as err:
                print(f'Except for item {item}')
                pprint(repr(err))
                not_entered_values.append(item)

        else:
            print(f'item not found {item["number"]}')
            not_entered_values.append(item)

    if not_entered_values:
        print(f'{"*" * 25} not_entered_values {"*" * 25}')
        pprint(not_entered_values)


def run_fast_scan_directory(directory: str, ext: list):
    """ Рекурсивный поиск файлов в папках по заданному пути (directory) с заданным расширением (ext)
    :param directory:
    :param ext:
    :return:
    """
    reports_sub_folders = []
    reports_files = []

    file: object

    for file in os.scandir(directory):
        if os.path.isdir(file) and getattr(file, 'path'):
            reports_sub_folders.append(file.path)

        if os.path.isfile(file) and getattr(file, 'name'):
            if os.path.splitext(file.name)[1].lower() not in ext:
                continue
            reports_files.append(file.path)

    directory: object
    for directory in list(reports_sub_folders):
        sf, file = run_fast_scan_directory(directory, ext)
        reports_sub_folders.extend(sf)
        reports_files.extend(file)

    return reports_sub_folders, reports_files


def main() -> None:
    """
    :return: None
    """
    report_extension = EXTENSION_DOCX
    report_files: list = []
    folders_with_data: list = []

    if _check_is_folder(path_folder=MAIN_FOLDER):
        folders_with_data = _collect_folder()

    for folder in folders_with_data:
        found_files = _collect_file(folder, report_extension=report_extension)

        for found_file in found_files:
            report_files.append({"file": found_file, "path": MAIN_FOLDER + "\\" + folder + "\\" + found_file})

    service = account_credentials(service_account_file=SERVICE_ACCOUNT_FILE)

    range_workers = ["Сотрудники!E9:E"]
    range_workers = get_data_from_google(service=service,
                                         spread_sheet_id=SPREAD_SHEET_ID,
                                         sheet_range=range_workers,
                                         name="sheet_values_workers")

    if report_files:
        js.write_json_file(data=report_files, name=PATH_TO_JSON + "report_data")
        load_reports(service=service,
                     reports_files=report_files,
                     reports_extension=report_extension,
                     range_workers=range_workers)


def start_load_report():
    """
    :return:
    """
    print("Starting load report")

    extension = ['.docx']
    report_files = []
    sub_folders, files = run_fast_scan_directory(MAIN_FOLDER, extension)

    for item in files:
        report_files.append({"file": item.split('\\')[-1], "path": item})

    if not report_files:
        print("!!No data for load!!")

    js.write_json_file(data=report_files, name=PATH_TO_JSON + "report_data")

    service = account_credentials(service_account_file=SERVICE_ACCOUNT_FILE)

    range_workers = ["Сотрудники!E9:E"]
    range_workers = get_data_from_google(service=service,
                                         spread_sheet_id=SPREAD_SHEET_ID,
                                         sheet_range=range_workers,
                                         name="sheet_values_workers")

    for ext in extension:
        data_for_load = _read_data_from_files(reports_files=report_files,
                                              reports_extension=str(ext),
                                              range_workers=range_workers)

        _load_reports(service=service,
                      range_workers=range_workers,
                      data_for_load=data_for_load)
    print("Finished loading reports")


if __name__ == '__main__':
    # start_load_report()

    main()
